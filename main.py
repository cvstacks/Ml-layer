import os
import requests
import tempfile
import pdfplumber
from docx import Document
from urllib.parse import urlparse

def download_resume(url):
    response = requests.get(url, stream=True)
    response.raise_for_status()

     # FIX: Strip query parameters (SAS tokens) before extracting extension
    parsed = urlparse(url)
    suffix = parsed.path.split('.')[-1]
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}")

    for chunk in response.iter_content(1024):
        temp_file.write(chunk)

    temp_file.close()
    return temp_file.name


def detect_file_type(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return "pdf"
    elif ext in [".docx", ".doc"]:
        return "docx"
    else:
        raise ValueError("Unsupported file type")


def extract_pdf(file_path):
    full_text = ""
    tables = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()

            if text:
                full_text += text + "\n"

            page_tables = page.extract_tables()
            for table in page_tables:
                tables.append(table)

    return { "text": full_text.strip(), "tables": tables, "hyperlinks": [] }


def extract_docx(file_path):

    doc = Document(file_path)

    full_text = "\n".join([p.text for p in doc.paragraphs])

    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])
        tables.append(table_data)

    return {
        "text": full_text.strip(),
        "tables": tables,
        "hyperlinks": []
    }


def extract_resume(file_path):

    file_type = detect_file_type(file_path)

    if file_type == "pdf":
        return extract_pdf(file_path)

    if file_type == "docx":
        return extract_docx(file_path)

    return None